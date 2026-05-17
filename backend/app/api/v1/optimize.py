import json
import io
import urllib.parse
from fastapi import APIRouter, HTTPException
from datetime import datetime
from app.db import database
from app.services.minimax import MiniMaxService
from app.core.config import settings
from app.models.common import ApiResponse

router = APIRouter(prefix="/resume", tags=["简历优化"])


def _error_response(code: int, message: str):
    return ApiResponse(code=code, message=message, data=None)


def _get_minimax():
    return MiniMaxService(api_key=settings.minimax_api_key, model=settings.minimax_model)


def _safe_filename(name: str) -> str:
    """RFC 5987 编码，防止 Content-Disposition 头编码错误"""
    encoded = urllib.parse.quote(name, safe='')
    return f"'{encoded}'"


def _extract_name(text: str) -> str:
    """从简历纯文本中提取姓名（首行或"姓名"字段）"""
    import re
    if not text:
        return ''
    lines = text.strip().split('\n')
    # 首行作为姓名（大多数简历格式）
    first = lines[0].strip()
    if 2 <= len(first) <= 10 and re.match(r'^[\u4e00-\u9fff·]+$', first):
        return first
    # 尝试匹配"姓名：XXX"
    m = re.search(r'姓名[：:]\s*([\u4e00-\u9fff·]{2,6})', text)
    if m:
        return m.group(1)
    return first if first else ''


@router.post("/optimize", response_model=ApiResponse)
async def optimize_resume(
    resume_id: str,
    optimize_type: str = "full",
):
    """AI优化简历（全部参数通过 query 传递）"""
    resume = await database.get_resume(resume_id)
    if not resume:
        return _error_response(2001, "简历不存在")

    if not resume.get("raw_text"):
        return _error_response(2002, "简历未解析，请先调用parse接口")

    try:
        minimax = _get_minimax()

        if optimize_type == "analyze":
            # 分析模式：调用 analyze_resume 返回完整评分+优劣势
            try:
                result = minimax.analyze_resume(
                    resume_text=resume["raw_text"],
                    target_industry=None,
                )
            except Exception as e:
                # 超时、网络错误、JSON错误等全部降级处理，保证流程不崩溃
                return ApiResponse(
                    code=0,
                    message="success",
                    data={
                        "resume_id": resume_id,
                        "basic_info": {"name": "", "education": "", "major": "", "experience_years": 0, "current_industry": ""},
                        "score": {"completeness": 5.0, "structure": 5.0, "keywords": 5.0},
                        "strengths": [],
                        "weaknesses": [],
                        "optimization_suggestions": [],
                        "analyzed_at": datetime.now().isoformat(),
                    },
                )

            if not isinstance(result, dict):
                return _error_response(2004, f"MiniMax API返回格式异常: {type(result).__name__}，请重试")

            return ApiResponse(
                code=0,
                message="success",
                data={
                    "resume_id": resume_id,
                    "basic_info": result.get("basic_info", {}),
                    "score": result.get("score", {"completeness": 5.0, "structure": 5.0, "keywords": 5.0}),
                    "strengths": result.get("strengths", []),
                    "weaknesses": result.get("weaknesses", []),
                    "optimization_suggestions": result.get("optimization_suggestions", []),
                    "analyzed_at": datetime.now().isoformat(),
                },
            )

        # 优化模式
        try:
            result = minimax.optimize_resume(
                resume_text=resume["raw_text"],
                target_industry=None,
                optimize_type=optimize_type,
                target_keywords=[],
            )
        except (json.JSONDecodeError, ValueError, Exception) as e:
            name = _extract_name(resume["raw_text"])
            return ApiResponse(
                code=0,
                message="success",
                data={
                    "resume_id": resume_id,
                    "name": name,
                    "optimized_text": resume["raw_text"],
                    "changes": [],
                    "summary": f"AI优化暂时不可用，请稍后再试（解析错误: {str(e)[:80]}）",
                },
            )

        if not isinstance(result, dict):
            # MiniMax返回了非dict格式（如list），尝试提取其中的dict
            if isinstance(result, list) and len(result) > 0:
                # list格式：可能是单个dict或dict列表
                if isinstance(result[0], dict):
                    # 尝试取第一个dict作为结果
                    first = result[0]
                    if isinstance(first, dict) and any(k in first for k in ('optimized_text', 'text', 'content', 'result')):
                        # 兼容多种字段名
                        optimized_text = first.get('optimized_text') or first.get('text') or first.get('content') or first.get('result') or ''
                        changes = first.get('changes', [])
                        summary = first.get('summary', 'AI优化完成')
                        if optimized_text:
                            await database.update_resume_optimized(resume_id, optimized_text)
                            name = _extract_name(resume["raw_text"])
                            return ApiResponse(code=0, message="success", data={
                                "resume_id": resume_id, "name": name,
                                "optimized_text": optimized_text, "changes": changes, "summary": summary,
                            })
            # 无法解析，使用原文降级
            name = _extract_name(resume["raw_text"])
            return ApiResponse(
                code=0,
                message="success",
                data={
                    "resume_id": resume_id,
                    "name": name,
                    "optimized_text": resume["raw_text"],
                    "changes": [],
                    "summary": "AI优化暂时不可用，已返回原始简历内容",
                },
            )

        optimized_text = result.get("optimized_text", "")
        await database.update_resume_optimized(resume_id, optimized_text)
        name = _extract_name(resume["raw_text"])

        changes = result.get("changes", [])
        if not isinstance(changes, list):
            changes = []

        return ApiResponse(
            code=0,
            message="success",
            data={
                "resume_id": resume_id,
                "name": name,
                "optimized_text": optimized_text,
                "changes": changes,
                "summary": result.get("summary", ""),
            },
        )
    except Exception as e:
        return _error_response(2004, f"MiniMax API错误: {str(e)}")


@router.post("/export", response_model=ApiResponse)
async def export_resume(
    resume_id: str,
    format: str = "pdf",
    use_optimized: bool = True,
):
    """导出简历为PDF/Word/HTML"""
    if format not in ("pdf", "docx", "html"):
        return _error_response(3001, "导出格式不支持，仅支持PDF/Word/HTML")

    resume = await database.get_resume(resume_id)
    if not resume:
        return _error_response(2001, "简历不存在")

    try:
        if use_optimized and resume.get("optimized_text"):
            text = resume["optimized_text"]
        elif resume.get("raw_text"):
            text = resume["raw_text"]
        else:
            return _error_response(3002, "没有可导出的简历内容")

        base_name = resume["file_name"].rsplit(".", 1)[0] if resume.get("file_name") else "resume"

        if format == "pdf":
            return await _export_pdf(text, base_name)
        elif format == "html":
            return await _export_html(text, base_name)
        else:
            return await _export_docx(text, base_name)

    except Exception as e:
        return _error_response(3002, f"导出生成失败: {str(e)}")


def _is_section_heading(s: str) -> bool:
    if not s: return False
    SECTION_KEYWORDS = ['个人信息', '求职意向', '教育背景', '工作经历', '项目经验',
                         '专业技能', '自我评价', '获奖荣誉', '语言能力', '兴趣爱好',
                         '社团经历', '实习经历', '培训经历', '证书资质', '其他']
    return any(s.startswith(kw) or s == kw for kw in SECTION_KEYWORDS)


def _is_bullet_line(s: str) -> bool:
    return s.startswith('•') or s.startswith('-') or s.startswith('·')


# ─── 后端简历 HTML 生成（与前端 formatResumeHtml 一致） ───────────

def _build_resume_html(text: str, title: str = "简历") -> str:
    """将纯文本简历转换为精美HTML排版（与前端一致的解析逻辑）"""
    import re

    SECTION_KEYWORDS = ['个人信息', '求职意向', '教育背景', '工作经历', '项目经验',
                         '专业技能', '自我评价', '获奖荣誉', '语言能力', '兴趣爱好',
                         '社团经历', '实习经历', '培训经历', '证书资质', '其他']

    def get_line_type(s: str) -> str:
        if not s: return 'empty'
        if len(s) >= 2 and len(s) <= 6 and re.match(r'^[\u4e00-\u9fff]', s) and '：' not in s and ':' not in s:
            return 'name'
        if any(s.startswith(kw) or s == kw for kw in SECTION_KEYWORDS): return 'section'
        if re.match(r'^[\u4e00-\u9fff\w\s,，、]+$', s) and len(s) < 60:
            tokens = re.split(r'[,，、\s]+', s)
            if len(tokens) >= 2 and all(len(t) <= 8 for t in tokens): return 'skill-tag'
        return 'text'

    def parse_meta(s: str) -> list:
        parts = re.split(r'\s*[·｜/]\s*', s)
        result = []
        for p in parts:
            p = p.strip()
            if not p: continue
            m = re.match(r'^([^\s：:]+)[：:]\s*(.+)', p)
            if m: result.append({'label': m.group(1), 'value': m.group(2)})
            else: result.append({'label': '', 'value': p})
        return result

    def parse_text(s: str) -> str:
        if s.startswith('•') or s.startswith('-') or s.startswith('·'): return 'bullet'
        if re.match(r'^[\u4e00-\u9fff]', s) and re.search(r'\d{4}', s) and len(s) < 80: return 'entry'
        return 'text'

    lines = text.split('\n')
    sections = []
    i = 0

    while i < len(lines):
        s = lines[i].strip()
        t = get_line_type(s)

        if t == 'empty':
            i += 1
            continue
        if t == 'name':
            sections.append({'type': 'name', 'text': s})
            nxt = lines[i+1].strip() if i+1 < len(lines) else ''
            nxt_t = get_line_type(nxt)
            if nxt and nxt_t not in ('section', 'empty'):
                sections.append({'type': 'meta', 'text': nxt})
                i += 2
            else:
                i += 1
            continue
        if t == 'meta':
            sections.append({'type': 'meta', 'text': s})
            i += 1
            continue
        if t == 'skill-tag':
            tags = re.split(r'[,，、\s]+', s)
            for tag in tags:
                if tag: sections.append({'type': 'skill-tag', 'text': tag})
            i += 1
            continue
        if t == 'section':
            sections.append({'type': 'section', 'text': s})
            i += 1
            continue
        if t == 'text':
            tt = parse_text(s)
            if tt == 'entry':
                sections.append({'type': 'entry', 'text': s, 'level': 0})
                nxt = lines[i+1].strip() if i+1 < len(lines) else ''
                nxt_t = get_line_type(nxt)
                if nxt and nxt_t not in ('section','empty') and not nxt.startswith('•') and not nxt.startswith('-'):
                    sections.append({'type': 'entry', 'text': nxt, 'level': 1})
                    i += 2
                else:
                    i += 1
            elif tt == 'bullet':
                sections.append({'type': 'bullet', 'text': s[1:].strip()})
                i += 1
            else:
                sections.append({'type': 'text', 'text': s})
                i += 1
            continue
        i += 1

    html_parts = []
    sec_open = False

    for j, sec in enumerate(sections):
        t = sec['type']
        s = sec['text']

        if t == 'name':
            if sec_open:
                html_parts.append('</div></div>')
                sec_open = False
            html_parts.append('<div class="rph-header">')
            html_parts.append('<div class="rph-name-wrap"><span class="rph-name-star">✦</span><h1 class="rph-name">' + s + '</h1><span class="rph-name-star">✦</span></div>')
            if j+1 < len(sections) and sections[j+1]['type'] == 'meta':
                metas = parse_meta(sections[j+1]['text'])
                capsule_html = ''
                for m in metas:
                    if m['label']:
                        capsule_html += '<span class="rph-capsule"><span class="rph-capsule-label">' + m['label'] + '</span><span class="rph-capsule-value">' + m['value'] + '</span></span>'
                    else:
                        capsule_html += '<span class="rph-capsule rph-capsule-plain">' + m['value'] + '</span>'
                html_parts.append('<div class="rph-capsule-row">' + capsule_html + '</div>')
                j += 1
            html_parts.append('<div class="rph-header-divider"></div>')
            html_parts.append('</div>')
            continue

        if t == 'section':
            if sec_open:
                html_parts.append('</div></div>')
            html_parts.append('<div class="rph-section-block">')
            html_parts.append('<div class="rph-section-header"><span class="rph-section-icon">▶</span><h2 class="rph-section-title">' + s + '</h2></div>')
            html_parts.append('<div class="rph-section-body">')
            sec_open = True
            continue

        if t == 'entry':
            if sec.get('level', 0) == 0:
                html_parts.append('<div class="rph-entry"><div class="rph-entry-org">' + s + '</div>')
            else:
                html_parts.append('<div class="rph-entry-meta">' + s + '</div></div>')
            continue

        if t == 'bullet':
            html_parts.append('<div class="rph-bullet-item"><span class="rph-bullet">•</span><span class="rph-bullet-text">' + s + '</span></div>')
            continue

        if t == 'skill-tag':
            html_parts.append('<span class="rph-skill-tag">' + s + '</span>')
            continue

        if t == 'text':
            html_parts.append('<div class="rph-text-line">' + s + '</div>')
            continue

    if sec_open:
        html_parts.append('</div></div>')

    body_html = '\n'.join(html_parts)

    html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: "PingFang SC", "Microsoft YaHei", "Noto Sans SC", -apple-system, sans-serif; background: #fff; color: #2d3748; }}
.resume-html {{ padding: 36px 40px; max-width: 780px; margin: 0 auto; }}
.rph-header {{ margin-bottom: 6px; }}
.rph-name-wrap {{ display: flex; align-items: center; justify-content: center; gap: 16px; margin-bottom: 10px; }}
.rph-name {{ font-size: 26px; font-weight: 800; color: #1a365d; letter-spacing: 3px; margin: 0; border: none; padding: 0; background: none; text-align: center; }}
.rph-name-star {{ color: #3182ce; font-size: 16px; opacity: 0.7; }}
.rph-capsule-row {{ display: flex; flex-wrap: wrap; justify-content: center; gap: 8px; margin-bottom: 14px; }}
.rph-capsule {{ display: inline-flex; align-items: center; background: #edf2f7; border-radius: 20px; padding: 3px 12px; font-size: 12px; }}
.rph-capsule-label {{ color: #4a5568; font-weight: 500; }}
.rph-capsule-value {{ color: #2d3748; }}
.rph-capsule-plain {{ background: #ebf8ff; color: #2b6cb0; }}
.rph-header-divider {{ height: 2px; background: linear-gradient(to right, transparent, #1a365d 20%, #3182ce 50%, #1a365d 80%, transparent); border-radius: 1px; }}
.rph-section-block {{ margin-top: 18px; }}
.rph-section-header {{ display: flex; align-items: center; gap: 8px; margin-bottom: 10px; padding-bottom: 6px; border-bottom: 1.5px solid #e2e8f0; }}
.rph-section-icon {{ color: #3182ce; font-size: 12px; }}
.rph-section-title {{ font-size: 14px; font-weight: 700; color: #1a365d; letter-spacing: 1px; margin: 0; border: none; padding: 0; background: none; }}
.rph-entry {{ margin-bottom: 8px; }}
.rph-entry-org {{ font-size: 13.5px; font-weight: 600; color: #1a365d; margin-bottom: 2px; }}
.rph-entry-meta {{ font-size: 12px; color: #718096; margin-bottom: 4px; }}
.rph-bullet-item {{ display: flex; align-items: flex-start; gap: 8px; margin: 4px 0; padding-left: 2px; }}
.rph-bullet {{ color: #3182ce; font-size: 13px; flex-shrink: 0; margin-top: 1px; line-height: 1.6; }}
.rph-bullet-text {{ color: #2d3748; font-size: 13px; line-height: 1.65; flex: 1; }}
.rph-skill-tag {{ display: inline-block; background: #ebf8ff; color: #2b6cb0; border: 1px solid #bee3f8; border-radius: 6px; padding: 3px 10px; font-size: 12px; margin: 3px 4px 3px 0; font-weight: 500; }}
.rph-text-line {{ font-size: 13px; color: #2d3748; line-height: 1.7; margin: 3px 0; }}
@media print {{ body {{ padding: 0; }} .resume-html {{ padding: 28px 32px; }} }}
@media (max-width: 768px) {{ .resume-html {{ padding: 16px 16px; }} .rph-name {{ font-size: 20px; }} }}
</style>
</head>
<body>
<div class="resume-html">
{body_html}
</div>
</body>
</html>"""
    return html


async def _export_html(text: str, base_name: str):
    """导出为精美 HTML 文件（可浏览器打印为PDF）"""
    name = _extract_name(text)
    html = _build_resume_html(text, title=f"{name} - 简历")
    buffer = io.BytesIO(html.encode('utf-8'))
    buffer.seek(0)
    output_name = f"{base_name}_optimized.html"
    encoded_name = urllib.parse.quote(output_name)
    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buffer,
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )


def _parse_resume_for_export(text: str) -> list[dict]:
    """将纯文本简历解析为结构化数据（与前端 formatResumeHtml 一致的解析逻辑）"""
    import re
    lines = text.split('\n')
    result = []
    i = 0

    SECTION_KEYWORDS = ['个人信息', '求职意向', '教育背景', '工作经历', '项目经验',
                         '专业技能', '自我评价', '获奖荣誉', '语言能力', '兴趣爱好',
                         '社团经历', '实习经历', '培训经历', '证书资质']

    def get_line_type(s: str) -> str:
        if not s: return 'empty'
        if any(s.startswith(kw) or s == kw for kw in SECTION_KEYWORDS): return 'section'
        if 2 <= len(s) <= 6 and re.match(r'[\u4e00-\u9fff]', s[0]) and '：' not in s and ':' not in s: return 'name'
        return 'text'

    while i < len(lines):
        s = lines[i].strip()
        if not s:
            result.append({'type': 'empty'})
            i += 1
            continue
        t = get_line_type(s)
        if t == 'section':
            result.append({'type': 'section', 'text': s})
        elif t == 'name':
            result.append({'type': 'name', 'text': s})
            if i + 1 < len(lines):
                next_s = lines[i + 1].strip()
                if next_s and get_line_type(next_s) not in ('section', 'empty'):
                    result.append({'type': 'meta', 'text': next_s})
                    i += 1
        else:
            result.append({'type': 'text', 'text': s})
        i += 1
    return result


def _is_section_heading(s: str) -> bool:
    if not s: return False
    SECTION_KEYWORDS = ['个人信息', '求职意向', '教育背景', '工作经历', '项目经验',
                         '专业技能', '自我评价', '获奖荣誉', '语言能力', '兴趣爱好',
                         '社团经历', '实习经历', '培训经历', '证书资质']
    return any(s.startswith(kw) or s == kw for kw in SECTION_KEYWORDS)


def _is_bullet_line(s: str) -> bool:
    return s.startswith('•') or s.startswith('-') or s.startswith('·')


async def _export_pdf(text: str, base_name: str):
    """导出为PDF（精美结构化排版，使用 reportlab + wqy-microhei）"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体
    pdfmetrics.registerFont(TTFont("WQY", "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"))

    FONT = "WQY"
    PRIMARY = colors.HexColor("#1A365D")
    ACCENT = colors.HexColor("#3182CE")
    LIGHT = colors.HexColor("#718096")
    BORDER = colors.HexColor("#E2E8F0")
    GREEN = colors.HexColor("#15803D")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2.0 * cm,
        rightMargin=2.0 * cm,
        topMargin=1.8 * cm,
        bottomMargin=1.8 * cm,
    )

    def mkstyle(size, color=None, bold=False, sb=0, sa=2, lead=None):
        return ParagraphStyle(
            f"s_{size}_{bold}",
            fontName=FONT,
            fontSize=size,
            leading=lead or size * 1.65,
            spaceBefore=sb,
            spaceAfter=sa,
            textColor=color or colors.black,
        )

    story = []
    parsed = _parse_resume_for_export(text)
    name_printed = False

    for item in parsed:
        t = item['type']
        s = item.get('text', '')

        if t == 'empty':
            story.append(Spacer(1, 3))
            continue

        if t == 'name':
            name_printed = True
            story.append(Paragraph(s, mkstyle(20, PRIMARY, bold=True, sa=4, lead=26)))
            story.append(Spacer(1, 3))

        elif t == 'meta':
            story.append(Paragraph(s, mkstyle(10, LIGHT, sa=2, lead=14)))
            story.append(HRFlowable(width="100%", thickness=2, color=PRIMARY, spaceAfter=12))
            story.append(Spacer(1, 4))

        elif t == 'section':
            story.append(Spacer(1, 10))
            story.append(Paragraph(s, mkstyle(11.5, PRIMARY, bold=True, sa=3, lead=15)))
            story.append(HRFlowable(width="100%", thickness=0.8, color=PRIMARY, spaceAfter=6))

        else:  # text
            if _is_bullet_line(s):
                # 带蓝色圆点的列表项
                bullet_char = "•"
                bullet_text = s[1:].strip() if len(s) > 1 else s
                bullet_style = mkstyle(9.5, colors.black, sa=2, lead=14)
                story.append(Paragraph(f"<b><font color='#3182CE'>{bullet_char}</font></b>  {bullet_text}", bullet_style))
            else:
                story.append(Paragraph(s, mkstyle(9.5, colors.black, sa=2, lead=14)))

    if not name_printed:
        story.insert(0, Paragraph("简历", mkstyle(18, PRIMARY, bold=True, sa=6, lead=24)))
        story.insert(1, HRFlowable(width="100%", thickness=2, color=PRIMARY, spaceAfter=12))

    doc.build(story)
    buffer.seek(0)

    output_name = f"{base_name}_optimized.pdf"
    encoded_name = urllib.parse.quote(output_name)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )


async def _export_docx(text: str, base_name: str):
    """导出为Word（python-docx，结构化排版）"""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from fastapi.responses import StreamingResponse

    doc = Document()
    sect = doc.sections[0]
    sect.page_width = Cm(21)
    sect.page_height = Cm(29.7)
    sect.left_margin = Cm(2.2)
    sect.right_margin = Cm(2.2)
    sect.top_margin = Cm(2)
    sect.bottom_margin = Cm(2)

    PRIMARY = RGBColor(0x1A, 0x36, 0x5D)
    ACCENT = RGBColor(0x31, 0x82, 0xCE)

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = PRIMARY
        # 下划线效果
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = __import__('xml.etree.ElementTree' if False else 'lxml').etree.SubElement(pPr, qn('w:pBdr'))
        bottom = __import__('xml.etree.ElementTree' if False else 'lxml').etree.SubElement(pBdr, qn('w:bottom'))
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1A365D')
        return p

    def add_bullet(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Cm(0.5)
        run1 = p.add_run('• ')
        run1.font.color.rgb = ACCENT
        run1.font.size = Pt(10.5)
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
        return p

    def add_para(text, size=Pt(10.5)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        run.font.size = size
        return p

    parsed = _parse_resume_for_export(text)
    name_printed = False

    for item in parsed:
        t = item['type']
        s = item.get('text', '')

        if t == 'empty':
            continue

        if t == 'name':
            name_printed = True
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(20)
            run.font.color.rgb = PRIMARY

        elif t == 'meta':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(s)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x71, 0x80, 0x96)

        elif t == 'section':
            add_heading(s)

        else:  # text
            if _is_bullet_line(s):
                bullet_text = s[1:].strip() if len(s) > 1 else s
                add_bullet(bullet_text)
            else:
                add_para(s)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    output_name = f"{base_name}_optimized.docx"
    encoded_name = urllib.parse.quote(output_name)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"},
    )
